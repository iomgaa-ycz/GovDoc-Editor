"""Ground truth 解析工具——从附件9和人类工作底稿中提取结构化数据。"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def parse_gold_checkpoints(path: Path) -> list[dict[str, str]]:
    """解析附件9（金标准审核点表）为结构化列表。

    复用 checkpoint_import 解析器，输出精简 dict 供 judge 对比。

    Args:
        path: 附件9 xls/xlsx 文件路径。

    Returns:
        每项包含 title, description, category 的 dict 列表。
    """
    from govdoc.parsers.checkpoint_import import parse_checkpoint_file

    checkpoints, _ = parse_checkpoint_file(path)
    return [
        {
            "title": cp.title,
            "description": cp.description,
            "category": cp.category.value,
        }
        for cp in checkpoints
    ]


def parse_human_workpaper(path: Path) -> dict[str, Any]:
    """解析人类撰写的工作底稿 docx 为结构化数据。

    人类工作底稿为固定模板（8行×2列表格），核心内容在 Row5「检查情况摘要」。

    Args:
        path: 人类工作底稿 .docx 文件路径。

    Returns:
        dict 包含：
          - project_name: 检查项目名
          - checked_unit: 被检查单位
          - summary_text: 检查情况摘要全文
          - findings_text: 从摘要中提取的具体发现列表（按分段/编号切分）
    """
    from docx import Document

    doc = Document(str(path))
    if not doc.tables:
        logger.warning("人类工作底稿无表格: %s", path)
        return {"project_name": "", "checked_unit": "", "summary_text": "", "findings_text": []}

    table = doc.tables[0]
    rows = table.rows

    checked_unit = rows[1].cells[1].text.strip() if len(rows) > 1 else ""
    project_name = rows[2].cells[1].text.strip() if len(rows) > 2 else ""
    summary_text = rows[5].cells[1].text.strip() if len(rows) > 5 else ""

    findings_text = [
        s.strip()
        for s in re.split(r"\n(?=\d+[、.]|招标文件)", summary_text)
        if s.strip() and not s.strip().startswith("根据")
    ]

    return {
        "project_name": project_name,
        "checked_unit": checked_unit,
        "summary_text": summary_text,
        "findings_text": findings_text,
    }
