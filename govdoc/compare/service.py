"""文档对比服务层。

该模块负责把 DOCX/PDF 转换为统一文本块，构建 N 文件匹配 payload，
生成高亮 DOCX 副本，并把 review.json 与下载文件写入运行时目录。
"""

from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
import json
import re
import shutil
import uuid

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from govdoc.compare.compare import find_nfile_exact_matches
from govdoc.compare.extractor import extract_docx_paragraphs, extract_markdown_paragraphs
from govdoc.schemas.compare import (
    CompareArtifacts,
    CompareBlockSegment,
    CompareCategory,
    CompareCategoryId,
    CompareDocument,
    CompareDocumentBlock,
    CompareDocuments,
    CompareDownloads,
    CompareFileMeta,
    CompareMatch,
    CompareOccurrence,
    CompareOccurrenceSegment,
    CompareResponse,
    CompareSummary,
)


CATEGORY_PRIORITY: dict[CompareCategoryId, int] = {
    "paragraph": 0,
    "sentence": 1,
}

CATEGORY_LABELS: dict[CompareCategoryId, str] = {
    "paragraph": "相同段落",
    "sentence": "相同句子",
}

CATEGORY_COLORS: dict[CompareCategoryId, str] = {
    "paragraph": "#f5b700",
    "sentence": "#12b5cb",
}

DOCX_HIGHLIGHT_COLORS: dict[CompareCategoryId, WD_COLOR_INDEX] = {
    "paragraph": WD_COLOR_INDEX.YELLOW,
    "sentence": WD_COLOR_INDEX.TURQUOISE,
}

SENTENCE_END_CHARS = {
    "。",
    "！",
    "？",
    "!",
    "?",
    "；",
    ";",
}

ALLOWED_SUFFIXES = {".docx", ".pdf"}
REVIEW_ID_RE = re.compile(r"^[a-f0-9]{12}$")


@dataclass(frozen=True)
class TextBlock:
    """文档中的一个段落块及其全文偏移范围。"""

    id: str
    index: int
    text: str
    start: int
    end: int


@dataclass(frozen=True)
class SentenceOccurrence:
    """一个句子在全文和段落中的出现位置。"""

    file_index: int
    text: str
    start: int
    end: int
    block_id: str
    block_index: int
    start_in_block: int
    end_in_block: int


@dataclass(frozen=True)
class MatchOccurrence:
    """一个匹配项在某个文件全文中的出现范围。"""

    file_index: int
    start: int
    end: int


@dataclass(frozen=True)
class MatchRecord:
    """服务层内部使用的匹配记录。"""

    id: str
    category: CompareCategoryId
    text: str
    file_occurrences: dict[int, list[MatchOccurrence]]


@dataclass(frozen=True)
class DocumentModel:
    """服务层内部使用的文档模型。"""

    file_index: int
    file_name: str
    suffix: str
    blocks: list[TextBlock]
    full_text: str


@dataclass(frozen=True)
class CompareDownload:
    """高亮 DOCX 下载文件信息。"""

    path: Path
    filename: str


def get_compare_root() -> Path:
    """返回文档对比运行时目录。"""
    from govdoc.storage.files import get_storage_root

    root = get_storage_root() / "compare"
    root.mkdir(parents=True, exist_ok=True)
    return root


def _prepare_output_root(output_root: Path | None) -> Path:
    """准备对比输出根目录，测试可通过参数覆盖。"""
    root = output_root or get_compare_root()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _create_review_dir(output_root: Path) -> tuple[str, Path]:
    """创建唯一 review 目录并返回 review_id。"""
    while True:
        review_id = uuid.uuid4().hex[:12]
        review_dir = output_root / review_id
        if not review_dir.exists():
            review_dir.mkdir(parents=True)
            return review_id, review_dir


def _sanitize_filename(name: str) -> str:
    """清理上传文件名，避免写入危险或不可移植字符。"""
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", name)
    return cleaned.strip("._") or "reviewed_document.docx"


def _resolve_min_segment_length(value: int | None) -> int:
    """解析连续公共片段最小长度。"""
    if value is not None:
        return value
    from govdoc.runtime import get_config

    return get_config().compare.min_segment_length


def _validate_file_count(count: int) -> None:
    """校验文件数量满足 N 文件对比要求和部署配置。"""
    if count < 2:
        raise ValueError("至少上传 2 份文件。")

    from govdoc.runtime import get_config

    max_files = get_config().compare.max_files
    if max_files is not None and count > max_files:
        raise ValueError(f"当前部署最多支持 {max_files} 份文件。")


def _ensure_supported_suffix(filename: str) -> str:
    """校验文件扩展名并返回小写扩展名。"""
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise ValueError(f"仅支持 DOCX 和 PDF 文件，收到: {suffix or '无扩展名'}")
    return suffix


def _extract_pdf_paragraphs(path: Path) -> list[str]:
    """通过 DocumentStore 缓存路径把 PDF 转换为 Markdown 段落。"""
    from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError

    from govdoc.runtime import get_config, get_document_store

    timeout = get_config().compare.pdf_timeout_s
    store = get_document_store()

    with ThreadPoolExecutor(max_workers=1) as pool:
        try:
            prepared_md = pool.submit(store.get_or_convert, path).result(timeout=timeout)
        except FuturesTimeoutError:
            raise RuntimeError(f"PDF 转换超时（{timeout}s）: {path.name}")

    markdown = prepared_md.read_text(encoding="utf-8")
    return extract_markdown_paragraphs(markdown)


def _build_document_model(file_index: int, file_name: str, path: Path) -> DocumentModel:
    """把 DOCX/PDF 段落转换为带全文偏移的内部文档模型。"""
    suffix = _ensure_supported_suffix(file_name)
    if suffix == ".docx":
        paragraphs = extract_docx_paragraphs(path)
    elif suffix == ".pdf":
        paragraphs = _extract_pdf_paragraphs(path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")

    blocks: list[TextBlock] = []
    cursor = 0
    for index, text in enumerate(paragraphs, start=1):
        start = cursor
        end = start + len(text)
        blocks.append(
            TextBlock(
                id=f"file-{file_index}-block-{index}",
                index=index,
                text=text,
                start=start,
                end=end,
            )
        )
        cursor = end + 1

    full_text = "\n".join(block.text for block in blocks)
    return DocumentModel(
        file_index=file_index,
        file_name=file_name,
        suffix=suffix,
        blocks=blocks,
        full_text=full_text,
    )


def _find_sentence_boundary(text: str, index: int) -> int | None:
    """检测 text[index] 是否为句子结束符，返回句子边界位置。"""
    current = text[index]

    if current in SENTENCE_END_CHARS:
        return index + 1

    if current == ".":
        lookahead = index + 1
        while lookahead < len(text) and text[lookahead].isspace():
            lookahead += 1
        if (
            lookahead >= len(text)
            or text[lookahead].isupper()
            or text[lookahead].isdigit()
            or text[lookahead] in {'"', "'"}
        ):
            return lookahead

    return None


def _trim_and_append_sentence(
    sentences: list[SentenceOccurrence],
    document: DocumentModel,
    block: TextBlock,
    start: int,
    end: int,
) -> None:
    """裁剪段落内 [start, end) 范围的首尾空白后追加到句子列表。"""
    text = block.text
    leading = start
    while leading < end and text[leading].isspace():
        leading += 1
    trailing = end
    while trailing > leading and text[trailing - 1].isspace():
        trailing -= 1

    if leading < trailing:
        sentences.append(
            SentenceOccurrence(
                file_index=document.file_index,
                text=text[leading:trailing],
                start=block.start + leading,
                end=block.start + trailing,
                block_id=block.id,
                block_index=block.index,
                start_in_block=leading,
                end_in_block=trailing,
            )
        )


def _iter_sentence_occurrences(document: DocumentModel) -> list[SentenceOccurrence]:
    """枚举文档中所有句子的全文位置和段落内位置。"""
    sentences: list[SentenceOccurrence] = []

    for block in document.blocks:
        text = block.text
        start = 0
        index = 0

        while index < len(text):
            boundary_end = _find_sentence_boundary(text, index)
            if boundary_end is not None:
                _trim_and_append_sentence(sentences, document, block, start, boundary_end)
                start = boundary_end
            index += 1

        _trim_and_append_sentence(sentences, document, block, start, len(text))

    return sentences


def _build_nfile_block_matches(documents: list[DocumentModel]) -> list[MatchRecord]:
    """在所有文件间查找完全相同的段落。"""
    documents_by_index = {doc.file_index: doc for doc in documents}
    all_items = {doc.file_index: [block.text for block in doc.blocks] for doc in documents}
    exact_matches = find_nfile_exact_matches(all_items)

    matches: list[MatchRecord] = []
    for exact in exact_matches:
        file_occurrences: dict[int, list[MatchOccurrence]] = {}
        for file_index, positions in exact.file_positions.items():
            doc = documents_by_index[file_index]
            occurrences: list[MatchOccurrence] = []
            for position in positions:
                block_index = position - 1
                if 0 <= block_index < len(doc.blocks):
                    block = doc.blocks[block_index]
                    occurrences.append(
                        MatchOccurrence(
                            file_index=file_index,
                            start=block.start,
                            end=block.end,
                        )
                    )
            if occurrences:
                file_occurrences[file_index] = occurrences

        if len(file_occurrences) >= 2:
            matches.append(
                MatchRecord(
                    id=f"paragraph-{len(matches) + 1:03d}",
                    category="paragraph",
                    text=exact.text,
                    file_occurrences=file_occurrences,
                )
            )

    return matches


def _sentence_covered_by_paragraph(
    doc: DocumentModel,
    sentence: SentenceOccurrence,
    paragraph_ranges_by_file: dict[int, list[tuple[int, int]]] | None,
) -> bool:
    """判断句子是否已被段落级匹配完全覆盖。"""
    if not paragraph_ranges_by_file or doc.file_index not in paragraph_ranges_by_file:
        return False
    return _is_covered_by_ranges(
        document=doc, start=sentence.start, end=sentence.end,
        ranges=paragraph_ranges_by_file[doc.file_index],
    )


def _build_nfile_sentence_matches(
    documents: list[DocumentModel],
    paragraph_ranges_by_file: dict[int, list[tuple[int, int]]] | None = None,
) -> list[MatchRecord]:
    """在所有文件间查找完全相同的句子（排除已被段落级覆盖的）。"""
    sentence_lookup: dict[str, dict[int, list[SentenceOccurrence]]] = defaultdict(
        lambda: defaultdict(list)
    )
    first_seen: dict[str, tuple[int, int]] = {}

    for doc in documents:
        for order, sentence in enumerate(_iter_sentence_occurrences(doc), start=1):
            if _sentence_covered_by_paragraph(doc, sentence, paragraph_ranges_by_file):
                continue
            sentence_lookup[sentence.text][doc.file_index].append(sentence)
            first_seen.setdefault(sentence.text, (doc.file_index, order))

    texts = [
        text
        for text, per_file_sentences in sentence_lookup.items()
        if len(per_file_sentences) >= 2
    ]
    texts.sort(key=lambda text: (first_seen[text][0], first_seen[text][1], text))

    matches: list[MatchRecord] = []
    for text in texts:
        file_occurrences = {
            file_index: [
                MatchOccurrence(
                    file_index=file_index,
                    start=sentence.start,
                    end=sentence.end,
                )
                for sentence in sentences
            ]
            for file_index, sentences in sentence_lookup[text].items()
            if sentences
        }
        if len(file_occurrences) >= 2:
            matches.append(
                MatchRecord(
                    id=f"sentence-{len(matches) + 1:03d}",
                    category="sentence",
                    text=text,
                    file_occurrences=file_occurrences,
                )
            )

    return matches


def _build_exact_ranges_by_file(matches: list[MatchRecord]) -> dict[int, list[tuple[int, int]]]:
    """把段落/句子匹配范围按文件聚合，供片段去重使用。"""
    ranges: dict[int, list[tuple[int, int]]] = defaultdict(list)
    for match in matches:
        for file_index, occurrences in match.file_occurrences.items():
            ranges[file_index].extend((occurrence.start, occurrence.end) for occurrence in occurrences)
    return {file_index: sorted(items) for file_index, items in ranges.items()}


def _is_covered_by_ranges(
    document: DocumentModel,
    start: int,
    end: int,
    ranges: list[tuple[int, int]],
) -> bool:
    """判断 [start, end) 中的非空白文本是否已被更高优先级匹配覆盖。"""
    if start >= end:
        return True

    current = start
    for range_start, range_end in ranges:
        if range_end <= current:
            continue
        if range_start >= end:
            break
        if range_start > current and document.full_text[current:range_start].strip():
            return False
        current = max(current, min(end, range_end))
        if current >= end:
            return True

    return not document.full_text[current:end].strip()


def _split_occurrence_by_blocks(
    blocks: list[TextBlock],
    occurrence: MatchOccurrence,
) -> list[CompareOccurrenceSegment]:
    """把全文匹配范围切分到对应段落块内。"""
    pieces: list[CompareOccurrenceSegment] = []

    for block in blocks:
        overlap_start = max(block.start, occurrence.start)
        overlap_end = min(block.end, occurrence.end)
        if overlap_start >= overlap_end:
            continue

        pieces.append(
            CompareOccurrenceSegment(
                file_index=occurrence.file_index,
                block_id=block.id,
                block_index=block.index,
                start=overlap_start - block.start,
                end=overlap_end - block.start,
            )
        )

    return pieces


def _build_annotations(
    document: DocumentModel,
    matches: list[MatchRecord],
) -> tuple[dict[str, list[dict]], dict[str, list[CompareOccurrence]]]:
    """生成段落渲染 annotation 和按 match 聚合的位置索引。"""
    block_annotations: dict[str, list[dict]] = defaultdict(list)
    match_segments: dict[str, list[CompareOccurrence]] = defaultdict(list)

    for match in matches:
        occurrences = match.file_occurrences.get(document.file_index, [])

        for occurrence in occurrences:
            pieces = _split_occurrence_by_blocks(document.blocks, occurrence)
            match_segments[match.id].append(
                CompareOccurrence(
                    file_index=document.file_index,
                    start=occurrence.start,
                    end=occurrence.end,
                    segments=pieces,
                )
            )
            for piece in pieces:
                block_annotations[piece.block_id].append(
                    {
                        "match_id": match.id,
                        "category": match.category,
                        "start": piece.start,
                        "end": piece.end,
                    }
                )

    return block_annotations, match_segments


def _pick_primary_match(match_ids: list[str], match_lookup: dict[str, MatchRecord]) -> str:
    """根据类别优先级选择重叠片段的主匹配。"""
    return sorted(
        match_ids,
        key=lambda item: (CATEGORY_PRIORITY[match_lookup[item].category], item),
    )[0]


def _segment_for_range(
    block: TextBlock,
    start: int,
    end: int,
    annotations: list[dict],
    match_lookup: dict[str, MatchRecord],
) -> CompareBlockSegment:
    """为段落内 [start, end) 构建单个渲染片段。"""
    active = [a for a in annotations if a["start"] < end and a["end"] > start]
    match_ids = sorted({str(a["match_id"]) for a in active})
    categories: list[CompareCategoryId] = sorted(
        {a["category"] for a in active},
        key=lambda item: CATEGORY_PRIORITY[item],
    )
    return CompareBlockSegment(
        text=block.text[start:end],
        match_ids=match_ids,
        categories=categories,
        primary_match_id=_pick_primary_match(match_ids, match_lookup) if match_ids else None,
    )


def _can_merge_segments(left: CompareBlockSegment, right: CompareBlockSegment) -> bool:
    """判断两个相邻片段是否可合并（匹配类型完全相同）。"""
    return (
        left.match_ids == right.match_ids
        and left.categories == right.categories
        and left.primary_match_id == right.primary_match_id
    )


def _build_block_segments(
    block: TextBlock,
    annotations: list[dict],
    match_lookup: dict[str, MatchRecord],
) -> list[CompareBlockSegment]:
    """按 annotation 边界把段落切成前端可渲染片段。"""
    if not annotations:
        return [
            CompareBlockSegment(
                text=block.text, match_ids=[], categories=[], primary_match_id=None,
            )
        ]

    boundaries = {0, len(block.text)}
    for annotation in annotations:
        boundaries.add(annotation["start"])
        boundaries.add(annotation["end"])

    ordered = sorted(boundaries)
    segments: list[CompareBlockSegment] = []

    for start, end in zip(ordered, ordered[1:]):
        if start == end:
            continue
        current = _segment_for_range(block, start, end, annotations, match_lookup)
        if segments and _can_merge_segments(segments[-1], current):
            segments[-1] = segments[-1].model_copy(update={"text": segments[-1].text + current.text})
        else:
            segments.append(current)

    return segments


def _serialize_document(
    document: DocumentModel,
    block_annotations: dict[str, list[dict]],
    match_lookup: dict[str, MatchRecord],
) -> CompareDocument:
    """序列化单个文档为前端展示结构。"""
    return CompareDocument(
        file_index=document.file_index,
        name=document.file_name,
        suffix=document.suffix,
        block_count=len(document.blocks),
        blocks=[
            CompareDocumentBlock(
                id=block.id,
                index=block.index,
                text=block.text,
                segments=_build_block_segments(
                    block=block,
                    annotations=block_annotations.get(block.id, []),
                    match_lookup=match_lookup,
                ),
            )
            for block in document.blocks
        ],
    )


def _serialize_matches(
    matches: list[MatchRecord],
    match_segments_by_file: dict[int, dict[str, list[CompareOccurrence]]],
) -> list[CompareMatch]:
    """序列化匹配记录为前端列表结构。"""
    serialized: list[CompareMatch] = []

    for match in matches:
        file_indices = sorted(match.file_occurrences)
        occurrences = {
            str(file_index): match_segments_by_file.get(file_index, {}).get(match.id, [])
            for file_index in file_indices
        }
        per_file_counts = {
            str(file_index): len(items) for file_index, items in occurrences.items()
        }
        occurrence_count = sum(per_file_counts.values())
        serialized.append(
            CompareMatch(
                id=match.id,
                category=match.category,
                label=CATEGORY_LABELS[match.category],
                color=CATEGORY_COLORS[match.category],
                text=match.text,
                length=len(match.text),
                file_indices=file_indices,
                occurrences=occurrences,
                per_file_counts=per_file_counts,
                file_count=len(file_indices),
                occurrence_count=occurrence_count,
            )
        )

    serialized.sort(
        key=lambda item: (
            CATEGORY_PRIORITY[item.category],
            -item.file_count,
            -item.length,
            item.id,
        )
    )
    return serialized


def _write_highlighted_review_copy(path: Path, document_payload: CompareDocument) -> None:
    """根据前端片段结构写出带高亮的 DOCX 副本。"""
    review_doc = Document()

    for block in document_payload.blocks:
        paragraph = review_doc.add_paragraph()
        for segment in block.segments:
            run = paragraph.add_run(segment.text)
            if segment.categories:
                color = DOCX_HIGHLIGHT_COLORS[segment.categories[0]]
                run.font.highlight_color = color

    review_doc.save(path)


def _build_categories() -> list[CompareCategory]:
    """返回前端类别筛选所需的固定类别列表。"""
    return [
        CompareCategory(
            id=category,
            label=CATEGORY_LABELS[category],
            color=CATEGORY_COLORS[category],
        )
        for category in ("paragraph", "sentence")
    ]


def _build_compare_response(
    review_id: str,
    review_dir: Path,
    stored_files: list[tuple[Path, str]],
    min_segment_length: int,
) -> CompareResponse:
    """构建完整对比响应并落盘 review.json 与下载文件。"""
    documents = [
        _build_document_model(file_index=index, file_name=name, path=path)
        for index, (path, name) in enumerate(stored_files)
    ]

    paragraph_matches = _build_nfile_block_matches(documents)
    paragraph_ranges = _build_exact_ranges_by_file(paragraph_matches)
    sentence_matches = _build_nfile_sentence_matches(
        documents, paragraph_ranges_by_file=paragraph_ranges,
    )
    all_matches = paragraph_matches + sentence_matches
    match_lookup = {match.id: match for match in all_matches}

    match_segments_by_file: dict[int, dict[str, list[CompareOccurrence]]] = {}
    document_payloads: list[CompareDocument] = []
    for document in documents:
        annotations, match_segments = _build_annotations(document=document, matches=all_matches)
        match_segments_by_file[document.file_index] = match_segments
        document_payloads.append(
            _serialize_document(
                document=document,
                block_annotations=annotations,
                match_lookup=match_lookup,
            )
        )

    serialized_matches = _serialize_matches(
        matches=all_matches,
        match_segments_by_file=match_segments_by_file,
    )

    download_names: dict[str, str] = {}
    downloads: dict[str, str] = {}
    for document_payload in document_payloads:
        key = str(document_payload.file_index)
        stem = Path(_sanitize_filename(document_payload.name)).stem
        download_names[key] = f"{stem}_reviewed.docx"
        download_path = review_dir / f"file_{document_payload.file_index}_reviewed.docx"
        _write_highlighted_review_copy(download_path, document_payload)
        downloads[key] = f"/api/v1/compare/{review_id}/download/{document_payload.file_index}"

    payload = CompareResponse(
        review_id=review_id,
        summary=CompareSummary(
            file_count=len(documents),
            files=[
                CompareFileMeta(
                    file_index=document.file_index,
                    name=document.file_name,
                    suffix=document.suffix,
                    paragraph_count=len(document.blocks),
                    block_count=len(document.blocks),
                )
                for document in documents
            ],
            common_paragraph_count=len(paragraph_matches),
            common_sentence_count=len(sentence_matches),
            common_segment_count=0,
            match_count=len(serialized_matches),
            min_segment_length=min_segment_length,
        ),
        documents=CompareDocuments(files=document_payloads),
        matches=serialized_matches,
        categories=_build_categories(),
        downloads=CompareDownloads(files=downloads),
        artifacts=CompareArtifacts(
            review_dir=str(review_dir),
            download_names=download_names,
        ),
    )

    (review_dir / "review.json").write_text(
        json.dumps(payload.model_dump(mode="json", by_alias=True), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return payload


def create_compare_bundle(
    files: list[tuple[Path, str]],
    output_root: Path | None = None,
    min_segment_length: int | None = None,
) -> CompareResponse:
    """从多个本地 DOCX/PDF 路径创建对比 review。"""
    _validate_file_count(len(files))
    resolved_min_segment_length = _resolve_min_segment_length(min_segment_length)
    root = _prepare_output_root(output_root)
    review_id, review_dir = _create_review_dir(root)

    uploads_dir = review_dir / "uploads"
    uploads_dir.mkdir(exist_ok=True)

    stored_files: list[tuple[Path, str]] = []
    for file_index, (source_path, source_name) in enumerate(files):
        _ensure_supported_suffix(source_name)
        stored_path = uploads_dir / f"file_{file_index}_{_sanitize_filename(source_name)}"
        shutil.copy2(source_path, stored_path)
        stored_files.append((stored_path, source_name))

    return _build_compare_response(
        review_id=review_id,
        review_dir=review_dir,
        stored_files=stored_files,
        min_segment_length=resolved_min_segment_length,
    )


def create_compare_bundle_from_bytes(
    files: list[tuple[bytes, str]],
    output_root: Path | None = None,
    min_segment_length: int | None = None,
) -> CompareResponse:
    """从上传字节内容创建 N 文件对比 review。"""
    _validate_file_count(len(files))
    resolved_min_segment_length = _resolve_min_segment_length(min_segment_length)
    root = _prepare_output_root(output_root)
    review_id, review_dir = _create_review_dir(root)

    uploads_dir = review_dir / "uploads"
    uploads_dir.mkdir(exist_ok=True)

    stored_files: list[tuple[Path, str]] = []
    for file_index, (content, source_name) in enumerate(files):
        _ensure_supported_suffix(source_name)
        stored_path = uploads_dir / f"file_{file_index}_{_sanitize_filename(source_name)}"
        stored_path.write_bytes(content)
        stored_files.append((stored_path, source_name))

    return _build_compare_response(
        review_id=review_id,
        review_dir=review_dir,
        stored_files=stored_files,
        min_segment_length=resolved_min_segment_length,
    )


def get_compare_download(
    review_id: str,
    file_index: int,
    output_root: Path | None = None,
) -> CompareDownload:
    """读取 review 元数据并返回指定文件的高亮 DOCX 下载信息。"""
    if not REVIEW_ID_RE.fullmatch(review_id) or file_index < 0:
        raise FileNotFoundError(review_id)

    root = _prepare_output_root(output_root)
    review_dir = root / review_id
    metadata_path = review_dir / "review.json"
    if not metadata_path.exists():
        raise FileNotFoundError(review_id)

    metadata = CompareResponse.model_validate_json(metadata_path.read_text(encoding="utf-8"))
    key = str(file_index)
    if key not in metadata.artifacts.download_names:
        raise FileNotFoundError(review_id)

    path = review_dir / f"file_{file_index}_reviewed.docx"
    if not path.exists():
        raise FileNotFoundError(review_id)

    return CompareDownload(path=path, filename=metadata.artifacts.download_names[key])
