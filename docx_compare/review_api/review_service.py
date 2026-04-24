from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
import json
import re
import shutil
import uuid

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from docx_common_text.compare import find_common_segments
from docx_common_text.extractor import extract_docx_paragraphs


CATEGORY_PRIORITY = {
    "paragraph": 0,
    "sentence": 1,
    "segment": 2,
}

CATEGORY_LABELS = {
    "paragraph": "\u76f8\u540c\u6bb5\u843d",
    "sentence": "\u76f8\u540c\u53e5\u5b50",
    "segment": "\u8fde\u7eed\u516c\u5171\u7247\u6bb5",
}

CATEGORY_COLORS = {
    "paragraph": "#f5b700",
    "sentence": "#12b5cb",
    "segment": "#ff7a59",
}

DOCX_HIGHLIGHT_COLORS = {
    "paragraph": WD_COLOR_INDEX.YELLOW,
    "sentence": WD_COLOR_INDEX.TURQUOISE,
    "segment": WD_COLOR_INDEX.PINK,
}

SENTENCE_END_CHARS = {
    "\u3002",
    "\uff01",
    "\uff1f",
    "!",
    "?",
    "\uff1b",
    ";",
}


@dataclass(frozen=True)
class TextBlock:
    id: str
    index: int
    text: str
    start: int
    end: int


@dataclass(frozen=True)
class SentenceOccurrence:
    text: str
    start: int
    end: int
    block_id: str
    block_index: int
    start_in_block: int
    end_in_block: int


@dataclass(frozen=True)
class MatchOccurrence:
    start: int
    end: int


@dataclass(frozen=True)
class MatchRecord:
    id: str
    category: str
    text: str
    first_occurrences: list[MatchOccurrence]
    second_occurrences: list[MatchOccurrence]


@dataclass(frozen=True)
class DocumentModel:
    side: str
    file_name: str
    blocks: list[TextBlock]
    full_text: str


def _sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", name)
    return cleaned.strip("._") or "reviewed_document"


def _build_document_model(side: str, file_name: str, path: Path) -> DocumentModel:
    paragraphs = extract_docx_paragraphs(path)
    blocks: list[TextBlock] = []
    cursor = 0

    for index, text in enumerate(paragraphs, start=1):
        start = cursor
        end = start + len(text)
        blocks.append(
            TextBlock(
                id=f"{side}-block-{index}",
                index=index,
                text=text,
                start=start,
                end=end,
            )
        )
        cursor = end + 1

    full_text = "\n".join(block.text for block in blocks)
    return DocumentModel(side=side, file_name=file_name, blocks=blocks, full_text=full_text)


def _iter_sentence_occurrences(document: DocumentModel) -> list[SentenceOccurrence]:
    sentences: list[SentenceOccurrence] = []

    for block in document.blocks:
        text = block.text
        start = 0
        index = 0

        while index < len(text):
            current = text[index]
            boundary_end: int | None = None

            if current in SENTENCE_END_CHARS:
                boundary_end = index + 1
            elif current == ".":
                lookahead = index + 1
                while lookahead < len(text) and text[lookahead].isspace():
                    lookahead += 1
                if (
                    lookahead >= len(text)
                    or text[lookahead].isupper()
                    or text[lookahead].isdigit()
                    or text[lookahead] in {'"', "'"}
                ):
                    boundary_end = lookahead

            if boundary_end is not None:
                leading = start
                while leading < boundary_end and text[leading].isspace():
                    leading += 1
                trailing = boundary_end
                while trailing > leading and text[trailing - 1].isspace():
                    trailing -= 1

                if leading < trailing:
                    sentences.append(
                        SentenceOccurrence(
                            text=text[leading:trailing],
                            start=block.start + leading,
                            end=block.start + trailing,
                            block_id=block.id,
                            block_index=block.index,
                            start_in_block=leading,
                            end_in_block=trailing,
                        )
                    )
                start = boundary_end

            index += 1

        leading = start
        while leading < len(text) and text[leading].isspace():
            leading += 1
        trailing = len(text)
        while trailing > leading and text[trailing - 1].isspace():
            trailing -= 1

        if leading < trailing:
            sentences.append(
                SentenceOccurrence(
                    text=text[leading:trailing],
                    start=block.start + leading,
                    end=block.start + trailing,
                    block_id=block.id,
                    block_index=block.index,
                    start_in_block=leading,
                    end_in_block=trailing,
                )
            )

    return sentences


def _build_exact_block_matches(
    first_document: DocumentModel,
    second_document: DocumentModel,
) -> list[MatchRecord]:
    first_lookup: dict[str, list[MatchOccurrence]] = defaultdict(list)
    second_lookup: dict[str, list[MatchOccurrence]] = defaultdict(list)

    for block in first_document.blocks:
        first_lookup[block.text].append(MatchOccurrence(start=block.start, end=block.end))
    for block in second_document.blocks:
        second_lookup[block.text].append(MatchOccurrence(start=block.start, end=block.end))

    matches: list[MatchRecord] = []
    seen: set[str] = set()

    for block in first_document.blocks:
        if block.text in second_lookup and block.text not in seen:
            seen.add(block.text)
            match_id = f"paragraph-{len(matches) + 1:03d}"
            matches.append(
                MatchRecord(
                    id=match_id,
                    category="paragraph",
                    text=block.text,
                    first_occurrences=first_lookup[block.text],
                    second_occurrences=second_lookup[block.text],
                )
            )

    return matches


def _build_exact_sentence_matches(
    first_document: DocumentModel,
    second_document: DocumentModel,
) -> list[MatchRecord]:
    first_sentences = _iter_sentence_occurrences(first_document)
    second_sentences = _iter_sentence_occurrences(second_document)

    first_lookup: dict[str, list[MatchOccurrence]] = defaultdict(list)
    second_lookup: dict[str, list[MatchOccurrence]] = defaultdict(list)

    for sentence in first_sentences:
        first_lookup[sentence.text].append(
            MatchOccurrence(start=sentence.start, end=sentence.end)
        )
    for sentence in second_sentences:
        second_lookup[sentence.text].append(
            MatchOccurrence(start=sentence.start, end=sentence.end)
        )

    matches: list[MatchRecord] = []
    seen: set[str] = set()

    for sentence in first_sentences:
        if sentence.text in second_lookup and sentence.text not in seen:
            seen.add(sentence.text)
            match_id = f"sentence-{len(matches) + 1:03d}"
            matches.append(
                MatchRecord(
                    id=match_id,
                    category="sentence",
                    text=sentence.text,
                    first_occurrences=first_lookup[sentence.text],
                    second_occurrences=second_lookup[sentence.text],
                )
            )

    return matches


def _build_segment_matches(
    first_document: DocumentModel,
    second_document: DocumentModel,
    min_segment_length: int,
    exact_matches: list[MatchRecord],
) -> list[MatchRecord]:
    exact_occurrence_ranges: set[tuple[str, int, int, int, int]] = set()

    for match in exact_matches:
        first_ranges = {(item.start, item.end) for item in match.first_occurrences}
        second_ranges = {(item.start, item.end) for item in match.second_occurrences}
        for first_start, first_end in first_ranges:
            for second_start, second_end in second_ranges:
                exact_occurrence_ranges.add(
                    (match.text, first_start, first_end, second_start, second_end)
                )

    segments = find_common_segments(
        first_text=first_document.full_text,
        second_text=second_document.full_text,
        min_length=min_segment_length,
    )

    matches: list[MatchRecord] = []

    for segment in segments:
        segment_key = (
            segment.text,
            segment.first_start,
            segment.first_end,
            segment.second_start,
            segment.second_end,
        )
        if segment_key in exact_occurrence_ranges:
            continue

        match_id = f"segment-{len(matches) + 1:03d}"
        matches.append(
            MatchRecord(
                id=match_id,
                category="segment",
                text=segment.text,
                first_occurrences=[
                    MatchOccurrence(
                        start=segment.first_start,
                        end=segment.first_end,
                    )
                ],
                second_occurrences=[
                    MatchOccurrence(
                        start=segment.second_start,
                        end=segment.second_end,
                    )
                ],
            )
        )

    return matches


def _split_occurrence_by_blocks(
    blocks: list[TextBlock],
    occurrence: MatchOccurrence,
) -> list[dict]:
    pieces: list[dict] = []

    for block in blocks:
        overlap_start = max(block.start, occurrence.start)
        overlap_end = min(block.end, occurrence.end)
        if overlap_start >= overlap_end:
            continue

        pieces.append(
            {
                "blockId": block.id,
                "blockIndex": block.index,
                "start": overlap_start - block.start,
                "end": overlap_end - block.start,
            }
        )

    return pieces


def _build_annotations(
    document: DocumentModel,
    matches: list[MatchRecord],
    side: str,
) -> tuple[dict[str, list[dict]], dict[str, list[dict]]]:
    block_annotations: dict[str, list[dict]] = defaultdict(list)
    match_segments: dict[str, list[dict]] = defaultdict(list)

    for match in matches:
        occurrences = (
            match.first_occurrences if side == "first" else match.second_occurrences
        )

        for occurrence in occurrences:
            pieces = _split_occurrence_by_blocks(document.blocks, occurrence)
            match_segments[match.id].append(
                {
                    "start": occurrence.start,
                    "end": occurrence.end,
                    "segments": pieces,
                }
            )
            for piece in pieces:
                block_annotations[piece["blockId"]].append(
                    {
                        "matchId": match.id,
                        "category": match.category,
                        "start": piece["start"],
                        "end": piece["end"],
                    }
                )

    return block_annotations, match_segments


def _pick_primary_match(match_ids: list[str], match_lookup: dict[str, MatchRecord]) -> str:
    return sorted(
        match_ids,
        key=lambda item: (CATEGORY_PRIORITY[match_lookup[item].category], item),
    )[0]


def _build_block_segments(
    block: TextBlock,
    annotations: list[dict],
    match_lookup: dict[str, MatchRecord],
) -> list[dict]:
    if not annotations:
        return [
            {
                "text": block.text,
                "matchIds": [],
                "categories": [],
                "primaryMatchId": None,
            }
        ]

    boundaries = {0, len(block.text)}
    for annotation in annotations:
        boundaries.add(annotation["start"])
        boundaries.add(annotation["end"])

    ordered = sorted(boundaries)
    segments: list[dict] = []

    for start, end in zip(ordered, ordered[1:]):
        if start == end:
            continue

        text = block.text[start:end]
        active = [
            annotation
            for annotation in annotations
            if annotation["start"] < end and annotation["end"] > start
        ]
        match_ids = sorted({annotation["matchId"] for annotation in active})
        categories = sorted(
            {annotation["category"] for annotation in active},
            key=lambda item: CATEGORY_PRIORITY[item],
        )
        primary_match_id = (
            _pick_primary_match(match_ids, match_lookup) if match_ids else None
        )

        current = {
            "text": text,
            "matchIds": match_ids,
            "categories": categories,
            "primaryMatchId": primary_match_id,
        }

        if (
            segments
            and segments[-1]["matchIds"] == current["matchIds"]
            and segments[-1]["categories"] == current["categories"]
            and segments[-1]["primaryMatchId"] == current["primaryMatchId"]
        ):
            segments[-1]["text"] += text
        else:
            segments.append(current)

    return segments


def _serialize_document(
    document: DocumentModel,
    block_annotations: dict[str, list[dict]],
    match_lookup: dict[str, MatchRecord],
) -> dict:
    return {
        "name": document.file_name,
        "blockCount": len(document.blocks),
        "blocks": [
            {
                "id": block.id,
                "index": block.index,
                "text": block.text,
                "segments": _build_block_segments(
                    block=block,
                    annotations=block_annotations.get(block.id, []),
                    match_lookup=match_lookup,
                ),
            }
            for block in document.blocks
        ],
    }


def _serialize_matches(
    matches: list[MatchRecord],
    first_match_segments: dict[str, list[dict]],
    second_match_segments: dict[str, list[dict]],
) -> list[dict]:
    serialized: list[dict] = []

    for match in matches:
        serialized.append(
            {
                "id": match.id,
                "category": match.category,
                "label": CATEGORY_LABELS[match.category],
                "color": CATEGORY_COLORS[match.category],
                "text": match.text,
                "length": len(match.text),
                "firstOccurrences": first_match_segments.get(match.id, []),
                "secondOccurrences": second_match_segments.get(match.id, []),
                "firstCount": len(match.first_occurrences),
                "secondCount": len(match.second_occurrences),
            }
        )

    serialized.sort(
        key=lambda item: (
            CATEGORY_PRIORITY[item["category"]],
            -item["length"],
            item["id"],
        )
    )
    return serialized


def _write_highlighted_review_copy(path: Path, document_payload: dict) -> None:
    review_doc = Document()

    for block in document_payload["blocks"]:
        paragraph = review_doc.add_paragraph()
        for segment in block["segments"]:
            run = paragraph.add_run(segment["text"])
            if segment["categories"]:
                color = DOCX_HIGHLIGHT_COLORS[segment["categories"][0]]
                run.font.highlight_color = color

    review_doc.save(path)


def create_review_bundle(
    first_path: Path,
    second_path: Path,
    output_root: Path,
    min_segment_length: int = 16,
    first_name: str | None = None,
    second_name: str | None = None,
) -> dict:
    review_id = uuid.uuid4().hex[:12]
    review_dir = output_root / review_id
    review_dir.mkdir(parents=True, exist_ok=True)

    first_source_name = first_name or first_path.name
    second_source_name = second_name or second_path.name

    uploads_dir = review_dir / "uploads"
    uploads_dir.mkdir(exist_ok=True)

    stored_first = uploads_dir / _sanitize_filename(first_source_name)
    stored_second = uploads_dir / _sanitize_filename(second_source_name)
    shutil.copy2(first_path, stored_first)
    shutil.copy2(second_path, stored_second)

    first_document = _build_document_model("first", first_source_name, stored_first)
    second_document = _build_document_model("second", second_source_name, stored_second)

    paragraph_matches = _build_exact_block_matches(first_document, second_document)
    sentence_matches = _build_exact_sentence_matches(first_document, second_document)
    segment_matches = _build_segment_matches(
        first_document=first_document,
        second_document=second_document,
        min_segment_length=min_segment_length,
        exact_matches=paragraph_matches + sentence_matches,
    )
    all_matches = paragraph_matches + sentence_matches + segment_matches
    match_lookup = {match.id: match for match in all_matches}

    first_annotations, first_match_segments = _build_annotations(
        document=first_document,
        matches=all_matches,
        side="first",
    )
    second_annotations, second_match_segments = _build_annotations(
        document=second_document,
        matches=all_matches,
        side="second",
    )

    first_payload = _serialize_document(
        document=first_document,
        block_annotations=first_annotations,
        match_lookup=match_lookup,
    )
    second_payload = _serialize_document(
        document=second_document,
        block_annotations=second_annotations,
        match_lookup=match_lookup,
    )
    serialized_matches = _serialize_matches(
        matches=all_matches,
        first_match_segments=first_match_segments,
        second_match_segments=second_match_segments,
    )

    downloads_dir = review_dir / "downloads"
    downloads_dir.mkdir(exist_ok=True)

    first_download_name = (
        f"{Path(_sanitize_filename(first_source_name)).stem}_reviewed.docx"
    )
    second_download_name = (
        f"{Path(_sanitize_filename(second_source_name)).stem}_reviewed.docx"
    )

    first_download_path = downloads_dir / "first_reviewed.docx"
    second_download_path = downloads_dir / "second_reviewed.docx"

    _write_highlighted_review_copy(first_download_path, first_payload)
    _write_highlighted_review_copy(second_download_path, second_payload)

    payload = {
        "reviewId": review_id,
        "summary": {
            "firstFileName": first_source_name,
            "secondFileName": second_source_name,
            "firstParagraphCount": len(first_document.blocks),
            "secondParagraphCount": len(second_document.blocks),
            "commonParagraphCount": len(paragraph_matches),
            "commonSentenceCount": len(sentence_matches),
            "commonSegmentCount": len(segment_matches),
            "matchCount": len(serialized_matches),
            "minSegmentLength": min_segment_length,
        },
        "documents": {
            "first": first_payload,
            "second": second_payload,
        },
        "matches": serialized_matches,
        "categories": [
            {
                "id": category,
                "label": CATEGORY_LABELS[category],
                "color": CATEGORY_COLORS[category],
            }
            for category in ("paragraph", "sentence", "segment")
        ],
        "downloads": {
            "first": f"/api/reviews/{review_id}/download/first",
            "second": f"/api/reviews/{review_id}/download/second",
        },
        "artifacts": {
            "reviewDir": str(review_dir),
            "firstDownloadName": first_download_name,
            "secondDownloadName": second_download_name,
        },
    }

    (review_dir / "review.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return payload
