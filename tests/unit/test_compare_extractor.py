"""文档对比 DOCX 提取器测试。"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from govdoc.compare.extractor import extract_docx_full_text, extract_docx_paragraphs


def _write_docx(path: Path) -> None:
    """写入包含正文段落和表格段落的 DOCX 测试文件。"""
    document = Document()
    document.add_paragraph("  第一段\xa0内容  ")
    document.add_paragraph("第二段")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "表格单元 A"
    table.cell(0, 1).text = "表格单元 B"
    document.save(path)


def test_extract_docx_paragraphs_includes_table_cells(tmp_path: Path) -> None:
    """提取器应按正文顺序返回段落和表格单元格文本。"""
    docx_path = tmp_path / "sample.docx"
    _write_docx(docx_path)

    paragraphs = extract_docx_paragraphs(docx_path)

    assert paragraphs == ["第一段 内容", "第二段", "表格单元 A", "表格单元 B"]


def test_extract_docx_full_text_joins_paragraphs(tmp_path: Path) -> None:
    """完整文本接口应使用换行符连接所有段落。"""
    docx_path = tmp_path / "sample.docx"
    _write_docx(docx_path)

    full_text = extract_docx_full_text(docx_path)

    assert full_text == "第一段 内容\n第二段\n表格单元 A\n表格单元 B"


def test_extract_empty_docx_returns_empty_list(tmp_path: Path) -> None:
    """空文档应返回空段落列表。"""
    docx_path = tmp_path / "empty.docx"
    document = Document()
    document.save(docx_path)

    paragraphs = extract_docx_paragraphs(docx_path)

    assert paragraphs == []


def test_extract_whitespace_only_paragraphs_are_skipped(tmp_path: Path) -> None:
    """仅含空白的段落应被过滤。"""
    docx_path = tmp_path / "whitespace.docx"
    document = Document()
    document.add_paragraph("   ")
    document.add_paragraph("有效段落")
    document.add_paragraph("\t\n")
    document.save(docx_path)

    paragraphs = extract_docx_paragraphs(docx_path)

    assert paragraphs == ["有效段落"]


def test_extract_invalid_zip_raises_exception(tmp_path: Path) -> None:
    """非法 ZIP 文件应抛出异常。"""
    bad_path = tmp_path / "bad.docx"
    bad_path.write_bytes(b"not a zip file")

    with pytest.raises(Exception):
        extract_docx_paragraphs(bad_path)


def test_extract_zip_without_document_xml_raises_value_error(tmp_path: Path) -> None:
    """缺少 word/document.xml 的 ZIP 应抛出 ValueError。"""
    import zipfile

    zip_path = tmp_path / "no_doc.docx"
    with zipfile.ZipFile(zip_path, "w") as zf:
        zf.writestr("content.txt", "hello")

    with pytest.raises(ValueError, match="not a valid DOCX"):
        extract_docx_paragraphs(zip_path)
